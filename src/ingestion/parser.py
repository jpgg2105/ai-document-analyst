"""Document parsing — extracts structured text from PDF, DOCX, and Markdown files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from src.api.middleware.logging import get_logger

logger = get_logger(__name__)

# Supported MIME ↔ extension mapping
SUPPORTED_TYPES: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "markdown",
    ".txt": "text",
}


@dataclass
class ParsedPage:
    """One logical page / section extracted from a document."""
    page_number: int
    text: str
    section: str = ""


@dataclass
class ParsedDocument:
    """Full parsed output of a single file."""
    filename: str
    file_type: str
    pages: list[ParsedPage]
    total_pages: int = 0

    def __post_init__(self) -> None:
        self.total_pages = len(self.pages)


# ---------------------------------------------------------------------------
# PDF parsing via PyMuPDF
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> list[ParsedPage]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    pages: list[ParsedPage] = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text.strip():
            pages.append(ParsedPage(page_number=page_num, text=text.strip()))
    doc.close()
    return pages


# ---------------------------------------------------------------------------
# DOCX parsing via python-docx
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> list[ParsedPage]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    current_section = ""
    pages: list[ParsedPage] = []
    buffer: list[str] = []
    page_num = 1

    for para in doc.paragraphs:
        # Detect headings to track sections
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Flush buffer as a page when we hit a new heading
            if buffer:
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        text="\n".join(buffer).strip(),
                        section=current_section,
                    )
                )
                page_num += 1
                buffer = []
            current_section = para.text.strip()

        if para.text.strip():
            buffer.append(para.text.strip())

    # Flush remaining text
    if buffer:
        pages.append(
            ParsedPage(
                page_number=page_num,
                text="\n".join(buffer).strip(),
                section=current_section,
            )
        )
    return pages


# ---------------------------------------------------------------------------
# Markdown / plain-text parsing
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def _parse_markdown(path: Path) -> list[ParsedPage]:
    raw = path.read_text(encoding="utf-8")
    sections = _HEADING_RE.split(raw)

    pages: list[ParsedPage] = []
    current_section = ""
    page_num = 1

    # Anything before the first heading
    preamble = sections[0].strip() if sections else ""
    if preamble:
        pages.append(ParsedPage(page_number=page_num, text=preamble))
        page_num += 1

    # Each heading produces 2 groups: the '#' chars and the heading text
    i = 1
    while i < len(sections) - 1:
        current_section = sections[i + 1].strip()
        body_idx = i + 2
        body = sections[body_idx].strip() if body_idx < len(sections) else ""
        if body:
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    text=body,
                    section=current_section,
                )
            )
            page_num += 1
        i += 3  # skip hashes, title, body

    return pages


def _parse_text(path: Path) -> list[ParsedPage]:
    raw = path.read_text(encoding="utf-8")
    return [ParsedPage(page_number=1, text=raw.strip())] if raw.strip() else []


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "markdown": _parse_markdown,
    "text": _parse_text,
}


def parse_document(path: Path) -> ParsedDocument:
    """Parse a document file and return structured pages/sections.

    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the path does not exist.
    """
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    file_type = SUPPORTED_TYPES.get(suffix)
    if file_type is None:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: {', '.join(SUPPORTED_TYPES.keys())}"
        )

    logger.info("parsing_document", filename=path.name, file_type=file_type)
    pages = _PARSERS[file_type](path)
    logger.info(
        "parsing_complete",
        filename=path.name,
        total_pages=len(pages),
    )
    return ParsedDocument(filename=path.name, file_type=file_type, pages=pages)
